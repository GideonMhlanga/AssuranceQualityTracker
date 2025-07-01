import streamlit as st
import pandas as pd
import numpy as np
import datetime as dt
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
import plotly.express as px
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from sqlalchemy import text, Date, Numeric
from database import get_conn
from datetime import date, datetime, timedelta

# Configuration
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
INVENTORY_CATEGORIES = {
    "chemicals": ["Reagent", "Solvent", "Buffer", "Indicator", "Acid", "Base", "Consumables"],
    "glassware": ["Volumetric", "Measuring", "Container", "Specialty"],
    "equipment": ["Instrument", "Device", "Tool", "Machine"]
}

GLASSWARE_STATUS_OPTIONS = ["OK", "Broken", "Chipped", "Missing"]
EQUIPMENT_STATUS_OPTIONS = ["Working", "Needs service", "Broken", "In repair", "Out for calibration"]

# ======================
# Utility Functions
# ======================

def validate_date(date_str):
    """Validate date in YYYY-MM-DD format with additional checks"""
    try:
        # First check the format
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        
        # Additional validation - date shouldn't be in the distant past
        if date_obj.year < 2000:
            return False
            
        return True
    except ValueError:
        return False

def format_date_for_display(date_str):
    """Convert YYYY-MM-DD to Month Year (e.g., September 2028) for display"""
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        return date_obj.strftime("%B %Y")
    except:
        return date_str

def calculate_category_stats(category):
    """Calculate statistics for a given inventory category"""
    df = st.session_state.lab_inventory[category].copy()
    stats = {
        'total': len(df),
        'low': 0,
        'expired': 0,
        'damaged': 0,
        'issues': 0
    }
    
    if category == "chemicals":
        stats['low'] = len(df[df['status'] == 'Low'])
        stats['expired'] = len(df[df['status'] == 'Expired'])
    elif category == "glassware":
        stats['damaged'] = len(df[df['status'] == 'Damaged'])
    elif category == "equipment":
        stats['issues'] = len(df[df['status'].str.contains('need|repair', case=False, na=False)])
    
    return stats

def style_inventory_rows(row):
    """Apply styling to inventory rows based on status"""
    styles = [''] * len(row)
    
    if 'status' in row.index:
        if row['status'] == 'Low':
            styles = ['background-color: #FFF3CD'] * len(row)
        elif row['status'] == 'Expired':
            styles = ['background-color: #F8D7DA'] * len(row)
        elif row['status'] == 'Damaged':
            styles = ['background-color: #F8D7DA'] * len(row)
        elif row['status'] in ['Needs service', 'In repair']:
            styles = ['background-color: #FFF3CD'] * len(row)

    numeric_cols = ['current', 'minimum', 'monthly']
    for col in numeric_cols:
        if col in row.index and pd.notna(row[col]) and str(row[col]).strip():
            try:
                row[col] = f"{float(row[col]):.2f}"
            except (ValueError, TypeError):
                row[col] = ""
    
    return styles

def calculate_months_of_stock(row):
    """Calculate months of remaining stock"""
    try:
        if (pd.notna(row.get('monthly')) and pd.notna(row.get('current'))) and \
           (str(row['monthly']).strip() and str(row['current']).strip()):
            monthly = float(row['monthly'])
            current = float(row['current'])
            if monthly > 0:
                return current / monthly
    except (TypeError, ValueError, KeyError):
        pass
    return None

def generate_unique_id(category):
    """Generate a unique ID for any inventory category"""
    prefix = {
        "chemicals": "CHEM",
        "glassware": "GLAS",
        "equipment": "EQUIP"
    }.get(category, "ITEM")
    
    conn = get_conn()
    try:
        result = conn.execute(
            text(f"SELECT MAX(id) FROM {category}")
        ).fetchone()
        
        max_id = result[0] if result[0] else f"{prefix}-000"
        num_part = int(max_id.split('-')[1]) + 1
        return f"{prefix}-{num_part:03d}"
    except Exception as e:
        st.error(f"Error generating ID: {e}")
        return f"{prefix}-{len(st.session_state.lab_inventory[category])+1:03d}"
    finally:
        if conn:
            conn.close()
def send_email_report(receiver_email, subject, body, attachment=None):
    """Send inventory report via email with improved error handling"""
    try:
        # Verify email configuration exists
        if "email" not in st.secrets:
            st.error("Email configuration not found in secrets")
            return False
            
        email_config = st.secrets["email"]
        
        # Validate required fields
        required_fields = ["sender", "password", "smtp_server", "smtp_port"]
        for field in required_fields:
            if field not in email_config:
                st.error(f"Missing required email configuration: {field}")
                return False
        
        # Create message container
        msg = MIMEMultipart()
        msg['From'] = email_config["sender"]
        msg['To'] = receiver_email
        msg['Subject'] = subject
        
        # Attach body
        msg.attach(MIMEText(body, 'plain'))
        
        # Attach file if provided
        if attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 
                          'attachment; filename="inventory_report.docx"')
            msg.attach(part)
        
        # Connect to server and send email
        try:
            # Use SMTP_SSL for port 465, regular SMTP for port 587
            if int(email_config["smtp_port"]) == 465:
                with smtplib.SMTP_SSL(
                    email_config["smtp_server"], 
                    int(email_config["smtp_port"])
                ) as server:
                    server.login(email_config["sender"], email_config["password"])
                    server.send_message(msg)
            else:
                with smtplib.SMTP(
                    email_config["smtp_server"], 
                    int(email_config["smtp_port"])
                ) as server:
                    server.starttls()
                    server.login(email_config["sender"], email_config["password"])
                    server.send_message(msg)
            
            return True
                
        except smtplib.SMTPAuthenticationError:
            st.error("Authentication failed. Please check your email and password.")
            return False
        except smtplib.SMTPException as e:
            st.error(f"SMTP error occurred: {str(e)}")
            return False
            
    except Exception as e:
        st.error(f"Unexpected error occurred: {str(e)}")
        return False

# ======================
# Database Functions
# ======================

def init_inventory_tables():
    """Initialize inventory tables in the database if they don't exist"""
    conn = None
    try:
        conn = get_conn()
        if not conn:
            st.error("Database connection failed")
            return False

        # Start a transaction
        with conn.begin():
            # Create all tables first
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS glassware (
                    id TEXT PRIMARY KEY,
                    item TEXT UNIQUE NOT NULL,
                    category TEXT NOT NULL,
                    total_quantity INTEGER NOT NULL DEFAULT 1,
                    broken_quantity INTEGER NOT NULL DEFAULT 0,
                    current INTEGER GENERATED ALWAYS AS (total_quantity - broken_quantity) STORED,
                    unit TEXT,
                    location TEXT,
                    comment TEXT,
                    status TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS equipment (
                    id TEXT PRIMARY KEY,
                    item TEXT UNIQUE NOT NULL,
                    category TEXT NOT NULL,
                    total_quantity INTEGER NOT NULL DEFAULT 1,
                    non_working_quantity INTEGER NOT NULL DEFAULT 0,
                    working_quantity INTEGER GENERATED ALWAYS AS (total_quantity - non_working_quantity) STORED,
                    unit TEXT,
                    location TEXT,
                    status TEXT,
                    last_calibration DATE,
                    comment TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS chemicals (
                    id TEXT PRIMARY KEY,
                    item TEXT UNIQUE NOT NULL,
                    category TEXT NOT NULL,
                    minimum NUMERIC(10,2),
                    current NUMERIC(10,2),
                    monthly NUMERIC(10,2),
                    unit TEXT,
                    expiry DATE,
                    location TEXT,
                    supplier TEXT,
                    comment TEXT,
                    status TEXT,
                    "Months Stock" NUMERIC(10,2),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            
            # Create sequences for ID generation
            for category in ["chemicals", "glassware", "equipment"]:
                conn.execute(text(f"""
                    CREATE SEQUENCE IF NOT EXISTS {category}_id_seq;
                    ALTER TABLE {category} ALTER COLUMN id SET DEFAULT 
                        CASE 
                            WHEN '{category}' = 'chemicals' THEN 'CHEM-' || LPAD(nextval('{category}_id_seq')::text, 3, '0')
                            WHEN '{category}' = 'glassware' THEN 'GLAS-' || LPAD(nextval('{category}_id_seq')::text, 3, '0')
                            WHEN '{category}' = 'equipment' THEN 'EQUIP-' || LPAD(nextval('{category}_id_seq')::text, 3, '0')
                        END;
                """))
                
                conn.execute(text(f"""
                    SELECT setval('{category}_id_seq', 
                        COALESCE((SELECT MAX(SUBSTRING(id FROM '[0-9]+$')::int) FROM {category}), 0) + 1);
                """))
            
            # Create triggers for each table
            for table in ["chemicals", "glassware", "equipment"]:
                conn.execute(text(f"""
                    CREATE OR REPLACE FUNCTION update_{table}_timestamp()
                    RETURNS TRIGGER AS $$
                    BEGIN
                        NEW.updated_at = CURRENT_TIMESTAMP;
                        RETURN NEW;
                    END;
                    $$ LANGUAGE plpgsql;
                """))
                
                conn.execute(text(f"""
                    DROP TRIGGER IF EXISTS trigger_update_{table}_timestamp ON {table};
                    CREATE TRIGGER trigger_update_{table}_timestamp
                    BEFORE UPDATE ON {table}
                    FOR EACH ROW EXECUTE FUNCTION update_{table}_timestamp();
                """))
        
        return True
        
    except Exception as e:
        st.error(f"Error initializing tables: {e}")
        return False
    finally:
        if conn:
            conn.close()
                
def load_from_db():
    """Load all inventory data from PostgreSQL"""
    inventory = {
        "chemicals": pd.DataFrame(),
        "glassware": pd.DataFrame(),
        "equipment": pd.DataFrame()
    }
    
    conn = get_conn()
    if conn:
        try:
            inventory["chemicals"] = pd.read_sql(text("SELECT * FROM chemicals"), conn)
            inventory["glassware"] = pd.read_sql(text("SELECT * FROM glassware"), conn)
            inventory["equipment"] = pd.read_sql(text("SELECT * FROM equipment"), conn)
        except Exception as e:
            st.error(f"Error loading inventory data: {e}")
        finally:
            conn.close()
    return inventory

def save_to_db(category, df):
    """Save a DataFrame back to the database with proper date and numeric handling
    
    Args:
        category (str): Inventory category ('chemicals', 'glassware', 'equipment')
        df (pd.DataFrame): DataFrame containing inventory data
        
    Raises:
        ValueError: If date formats are invalid or data validation fails
    """
    conn = None
    try:
        conn = get_conn()
        if not conn:
            st.error("Database connection failed")
            return False

        df_to_save = df.copy()
        
        # =============================================
        # Enhanced Date Handling
        # =============================================
        date_columns = {
            'chemicals': ['expiry'],
            'equipment': ['last_calibration']
        }.get(category, [])
        
        for col in date_columns:
            if col in df_to_save.columns:
                # Convert to datetime and validate format
                df_to_save[col] = pd.to_datetime(
                    df_to_save[col], 
                    format='%Y-%m-%d', 
                    errors='coerce'
                )
                
                # Check for invalid dates
                invalid_dates = df_to_save[df_to_save[col].isna()]
                if not invalid_dates.empty:
                    st.error(f"Invalid {col} dates found in rows: {invalid_dates.index.tolist()}")
                    st.error("Please use YYYY-MM-DD format for all dates")
                    return False
                
                # Convert to date strings for database storage
                df_to_save[col] = df_to_save[col].dt.strftime('%Y-%m-%d')

        # =============================================
        # Numeric Field Handling
        # =============================================
        if category == "chemicals":
            numeric_cols = ['minimum', 'current', 'monthly', 'Months Stock']
            for col in numeric_cols:
                if col in df_to_save.columns:
                    df_to_save[col] = (
                        pd.to_numeric(df_to_save[col], errors='coerce')
                        .round(2)
                        .replace({np.nan: None})
                    )
            
            # Recalculate months of stock if needed
            if 'Months Stock' not in df_to_save.columns:
                df_to_save['Months Stock'] = df_to_save.apply(
                    lambda row: round(calculate_months_of_stock(row), 2), 
                    axis=1
                )

        # =============================================
        # Database Operations
        # =============================================
        with conn.begin():  # This handles transactions automatically
            # Clear existing data
            conn.execute(text(f"TRUNCATE TABLE {category} RESTART IDENTITY CASCADE"))
            
            # Insert new data with explicit data types
            dtype_mapping = {
                'chemicals': {
                    'expiry': Date,
                    'Months Stock': Numeric(10,2),
                    'minimum': Numeric(10,2),
                    'current': Numeric(10,2),
                    'monthly': Numeric(10,2)
                },
                'equipment': {
                    'last_calibration': Date
                }
            }
            
            df_to_save.to_sql(
                name=category,
                con=conn,
                if_exists='append',
                index=False,
                dtype=dtype_mapping.get(category, None)
            )
            
        return True
        
    except ValueError as ve:
        st.error(f"Data validation error: {str(ve)}")
        if conn:
            conn.rollback()
        return False
    except Exception as e:
        st.error(f"Database operation failed: {str(e)}")
        if conn:
            conn.rollback()
        return False
    finally:
        if conn:
            conn.close()

def insert_item(category, item_data):
    """Insert a new item into the specified category with proper date handling"""
    conn = None
    try:
        conn = get_conn()
        if not conn:
            st.error("Database connection failed")
            return None

        if category == "chemicals" and "Months Stock" not in item_data:
            item_data["Months Stock"] = calculate_months_of_stock(item_data)

        # Remove 'current' column for glassware since it's generated by the database
        if category in ["glassware", "equipment"]  and "current" in item_data:
            item_data.pop("current", None)

        columns = []
        placeholders = []
        params = {}
        
        for col, val in item_data.items():
            if ' ' in col:
                columns.append(f'"{col}"')
            else:
                columns.append(col)
            
            param_name = col.replace(' ', '_')
            params[param_name] = val
            placeholders.append(f':{param_name}')

        query = text(f"""
            INSERT INTO {category} ({', '.join(columns)})
            VALUES ({', '.join(placeholders)})
            RETURNING *
        """)

        result = conn.execute(query, params).fetchone()
        conn.commit()
        return result

    except Exception as e:
        if conn:
            conn.rollback()
        st.error(f"Database operation failed: {e}")
        return None
    finally:
        if conn:
            conn.close()

# ======================
# Core Functions
# ======================

def initialize_inventory_data():
    """Initialize inventory data structure with PostgreSQL"""
    try:
        init_inventory_tables()
        st.session_state.lab_inventory = load_from_db()
        
        if all(df.empty for df in st.session_state.lab_inventory.values()):
            st.warning("No inventory data found - loading sample data...")
            load_sample_data()
            
            for category, df in st.session_state.lab_inventory.items():
                save_to_db(category, df)
            
            st.session_state.lab_inventory = load_from_db()
        
        check_restock_status()
        
    except Exception as e:
        st.error(f"Error initializing inventory data: {e}")
        st.warning("Falling back to sample data...")
        load_sample_data()

def load_sample_data():
    """Load sample inventory data with YYYY-MM-DD date formats"""
    sample_chemicals = [
        {
            "id": "CHEM-001", 
            "item": "DPD#4 tablets", 
            "category": "Reagent", 
            "minimum": 750.0, 
            "current": 410.0, 
            "monthly": 125.0, 
            "unit": "tablets", 
            "expiry": "2028-09-01", 
            "location": "Shelf A1", 
            "supplier": "Sigma-Aldrich", 
            "comment": "", 
            "status": "Low",
            "Months Stock": 3.28
        },
        {
            "id": "CHEM-002", 
            "item": "Sodium Thiosulfate", 
            "category": "Reagent", 
            "minimum": 500.0, 
            "current": 1200.0, 
            "monthly": 250.0, 
            "unit": "grams", 
            "expiry": "2026-05-15", 
            "location": "Shelf B2", 
            "supplier": "Fisher Scientific", 
            "comment": "Keep in amber bottle", 
            "status": "OK",
            "Months Stock": 4.8
        },
        {
            "id": "CHEM-003", 
            "item": "Hydrochloric Acid", 
            "category": "Acid", 
            "minimum": 1000.0, 
            "current": 2500.0, 
            "monthly": 500.0, 
            "unit": "mL", 
            "expiry": "2029-12-31", 
            "location": "Acid Cabinet", 
            "supplier": "VWR", 
            "comment": "Concentrated - handle with care", 
            "status": "OK",
            "Months Stock": 5.0
        },
        {
            "id": "CHEM-004", 
            "item": "Phenolphthalein", 
            "category": "Indicator", 
            "minimum": 50.0, 
            "current": 25.0, 
            "monthly": 10.0, 
            "unit": "mL", 
            "expiry": "2023-11-30", 
            "location": "Shelf C3", 
            "supplier": "Sigma-Aldrich", 
            "comment": "Expired - needs disposal", 
            "status": "Expired",
            "Months Stock": 2.5
        }
    ]
    
    sample_glassware = [
        {
            "id": "GLAS-001", 
            "item": "100 mL Volumetric Flask", 
            "category": "Volumetric", 
            "current": 12, 
            "unit": "pieces", 
            "location": "Cabinet 1", 
            "comment": "Class A", 
            "status": "OK"
        },
        {
            "id": "GLAS-002", 
            "item": "50 mL Burette", 
            "category": "Volumetric", 
            "current": 6, 
            "unit": "pieces", 
            "location": "Cabinet 2", 
            "comment": "With PTFE stopcock", 
            "status": "Damaged"
        },
        {
            "id": "GLAS-003", 
            "item": "250 mL Beaker", 
            "category": "Measuring", 
            "current": 24, 
            "unit": "pieces", 
            "location": "Shelf D4", 
            "comment": "", 
            "status": "OK"
        }
    ]
    
    sample_equipment = [
        {
            "id": "EQUIP-001", 
            "item": "pH Meter", 
            "category": "Instrument", 
            "current": 3, 
            "unit": "units", 
            "location": "Bench 1", 
            "status": "Working", 
            "last_calibration": "2023-10-15", 
            "comment": "Calibration due in 3 months"
        },
        {
            "id": "EQUIP-002", 
            "item": "Analytical Balance", 
            "category": "Instrument", 
            "current": 2, 
            "unit": "units", 
            "location": "Balance Room", 
            "status": "Needs service", 
            "last_calibration": "2023-09-01", 
            "comment": "Drifting - needs technician"
        },
        {
            "id": "EQUIP-003", 
            "item": "Hot Plate", 
            "category": "Device", 
            "current": 5, 
            "unit": "units", 
            "location": "Bench 3", 
            "status": "Working", 
            "last_calibration": "2023-08-20", 
            "comment": ""
        }
    ]
    
    st.session_state.lab_inventory = {
        "chemicals": pd.DataFrame(sample_chemicals),
        "glassware": pd.DataFrame(sample_glassware),
        "equipment": pd.DataFrame(sample_equipment)
    }

def check_restock_status():
    """Check and update status for all inventory items"""
    today = datetime.now().date()
    
    # Update chemicals status
    if "chemicals" in st.session_state.lab_inventory:
        df = st.session_state.lab_inventory["chemicals"].copy()
        
        # Convert expiry to datetime if it's not already
        if 'expiry' in df.columns and df['expiry'].dtype == object:
            df['expiry'] = pd.to_datetime(df['expiry'], errors='coerce').dt.date
        
        # Calculate status based on conditions
        conditions = [
            (df['current'] <= df['minimum'] * 0.2) & (df['current'] > 0),
            df['current'] <= 0,
            (df['expiry'] < today) if 'expiry' in df.columns else False,
            (df['current'] <= df['minimum'] * 0.5) & (df['current'] > df['minimum'] * 0.2)
        ]
        choices = ["Critical", "Out of Stock", "Expired", "Low"]
        
        df['status'] = np.select(conditions, choices, default="OK")
        
        # Calculate months of stock
        if 'monthly' in df.columns and 'current' in df.columns:
            df['Months Stock'] = df.apply(
                lambda row: round(calculate_months_of_stock(row), 2) if pd.notna(row['monthly']) and float(row['monthly']) > 0 else None,
                axis=1
            )
        
        st.session_state.lab_inventory["chemicals"] = df
    
    # Update glassware status
    if "glassware" in st.session_state.lab_inventory:
        df = st.session_state.lab_inventory["glassware"].copy()
        if 'status' not in df.columns:
            df['status'] = "OK"
        st.session_state.lab_inventory["glassware"] = df
    
    # Update equipment status
    if "equipment" in st.session_state.lab_inventory:
        df = st.session_state.lab_inventory["equipment"].copy()
        if 'status' not in df.columns:
            df['status'] = "Working"
        st.session_state.lab_inventory["equipment"] = df

# ======================
# Inventory Management
# ======================

def manage_chemicals():
    """Chemical inventory management"""
    df = st.session_state.lab_inventory["chemicals"].copy()
    
    with st.expander("➕ Add New Chemical", expanded=False):
        with st.form("add_chemical_form"):
            cols = st.columns([1, 2, 1])
            with cols[0]:
                new_id = generate_unique_id("chemicals")
                st.text_input("ID", value=new_id, disabled=True)
                item_name = st.text_input("Item Name", key="new_chem_name")
                category = st.selectbox("Category", INVENTORY_CATEGORIES["chemicals"])
            with cols[1]:
                minimum_stock = st.number_input("Minimum Stock", min_value=0.0, step=0.1, format="%.2f")
                current_stock = st.number_input("Current Stock", min_value=0.0, step=0.1, format="%.2f")
                monthly_usage = st.number_input("Monthly Usage", min_value=0.0, step=0.1, format="%.2f")
                unit = st.selectbox("Unit", ["g", "kg", "mL", "L", "tablets", "bottles", "sachets", "units"])
            with cols[2]:
                expiry_date = st.date_input(
                    "Expiry Date",
                    min_value=date(2000, 1, 1),
                    format="YYYY-MM-DD"
                ).strftime("%Y-%m-%d")  # Convert to string format
                location = st.text_input("Storage Location")
                supplier = st.text_input("Supplier")
                comment = st.text_input("Comments")
            
            if st.form_submit_button("Add Chemical"):
                # Remove date validation since we're using date_input
                existing_items = df['item'].str.lower().tolist()
                if item_name.lower() in existing_items:
                    st.error(f"A chemical with the name '{item_name}' already exists in Inventory!")
                else:
                    new_item = {
                        "id": new_id,
                        "item": item_name,
                        "category": category,
                        "minimum": minimum_stock,
                        "current": current_stock,
                        "monthly": monthly_usage,
                        "unit": unit,
                        "expiry": expiry_date,
                        "location": location,
                        "supplier": supplier,
                        "comment": comment,
                        "status": "OK"
                    }
                    
                    # Calculate months of stock
                    if monthly_usage > 0:
                        new_item["Months Stock"] = round(current_stock / monthly_usage, 2)
                    
                    # Double-check ID doesn't exist
                    conn = get_conn()
                    if conn:
                        existing_id = conn.execute(
                            text("SELECT id FROM chemicals WHERE id = :id"),
                            {"id": new_id}
                        ).fetchone()
                        
                        if existing_id:
                            st.error(f"ID {new_id} already exists! Generating new ID...")
                            new_id = generate_unique_id("chemicals")
                            new_item["id"] = new_id
                    
                    result = insert_item("chemicals", new_item)
                    if result:
                        st.session_state.lab_inventory["chemicals"] = pd.read_sql(
                            text("SELECT * FROM chemicals"), 
                            get_conn()
                        )
                        st.success(f"Added {item_name} to inventory")
                        st.rerun()
    
    st.subheader("Current Chemical Inventory")

    # Define the styling function
    def highlight_status(row):
        styles = [''] * len(row)
        if row['status'] == 'Expired':
            styles = ['background-color: #F8D7DA'] * len(row)  # Light red for expired
        elif row['status'] == 'Low':
            styles = ['background-color: #FFF3CD'] * len(row)  # Light yellow for low stock
        elif row['status'] == 'Critical':
            styles = ['background-color: #FFCCCB'] * len(row)  # Light red-orange for critical
        return styles
    
     # Display the styled dataframe
    st.dataframe(
        df.style.apply(highlight_status, axis=1),
        use_container_width=True,
        height=600,
        column_order=['id', 'item', 'category', 'current', 'minimum', 'monthly', 'Months Stock', 'status', 'expiry']
    )
    
    edited_df = st.data_editor(
        df,
        num_rows="dynamic",
        use_container_width=True,
        disabled=["id", "Months Stock"],
        column_config={
            "minimum": st.column_config.NumberColumn(format="%.2f"),
            "current": st.column_config.NumberColumn(format="%.2f"),
            "monthly": st.column_config.NumberColumn(format="%.2f"),
            "Months Stock": st.column_config.NumberColumn(format="%.2f"),
            "expiry": st.column_config.DateColumn(
            "Expiry Date",
            format="YYYY-MM-DD",
            min_value=date(2000, 1, 1),
            help="Select expiry date from calendar"
        ),
            "status": st.column_config.SelectboxColumn(
                "Status",
                options=["OK", "Low", "Critical", "Out of Stock", "Expired"],
                required=True
            )
        }
    )
    
    if st.button("Save Chemical Changes"):
        # Validate dates
        if 'expiry' in edited_df.columns:
            edited_df['expiry'] = pd.to_datetime(
                edited_df['expiry'], 
                format='%Y-%m-%d', 
                errors='coerce'
            )
            invalid_dates = edited_df[edited_df['expiry'].isna()]
            if not invalid_dates.empty:
                st.error("Some expiry dates are invalid. Please use YYYY-MM-DD format.")
                st.stop()
        
        # Recalculate months of stock
        if 'monthly' in edited_df.columns and 'current' in edited_df.columns:
            edited_df['Months Stock'] = edited_df.apply(
                lambda row: round(calculate_months_of_stock(row), 2), 
                axis=1
            )
        
        
        save_to_db("chemicals", edited_df)
        st.session_state.lab_inventory["chemicals"] = pd.read_sql(
            text("SELECT * FROM chemicals"), 
            get_conn()
        )
        st.success("Changes saved successfully!")
        st.rerun()

def manage_glassware():
    """Glassware inventory management with broken items tracking"""
    df = st.session_state.lab_inventory["glassware"].copy()
    
    with st.expander("➕ Add New Glassware", expanded=False):
        with st.form("add_glassware_form"):
            cols = st.columns([1, 2, 1])
            with cols[0]:
                new_id = generate_unique_id("glassware")
                st.text_input("ID", value=new_id, disabled=True)
                item_name = st.text_input("Item Name", key="new_glass_name")
                category = st.selectbox("Category", INVENTORY_CATEGORIES["glassware"])
            with cols[1]:
                total_quantity = st.number_input("Total Quantity", min_value=0, step=1, value=1)
                broken_quantity = st.number_input("Broken Quantity", min_value=0, step=1, value=0)
                usable_quantity = total_quantity - broken_quantity
                st.text_input("Usable Quantity", value=usable_quantity, disabled=True)
                unit = st.selectbox("Unit", ["pieces", "sets", "units"])
            with cols[2]:
                location = st.text_input("Storage Location")
                status = st.selectbox("Status", GLASSWARE_STATUS_OPTIONS)
                comment = st.text_input("Comments")
            
            if st.form_submit_button("Add Glassware"):
                existing_items = df['item'].str.lower().tolist()
                if item_name.lower() in existing_items:
                    st.error(f"Glassware with the name '{item_name}' already exists in Inventory")
                else:
                    new_item = {
                        "id": new_id,
                        "item": item_name,
                        "category": category,
                        "total_quantity": total_quantity,
                        "broken_quantity": broken_quantity,
                        # Remove "current": usable_quantity since it's generated by the database
                        "unit": unit,
                        "location": location,
                        "status": status if broken_quantity > 0 else "OK",
                        "comment": comment
                    }
                    
                    result = insert_item("glassware", new_item)
                    if result:
                        st.session_state.lab_inventory["glassware"] = pd.read_sql(
                            text("SELECT * FROM glassware"), 
                            get_conn()
                        )
                        st.success(f"Added {item_name} to inventory")
                        st.rerun()
    
    st.subheader("Current Glassware Inventory")
    
    # Add filtering options
    col1, col2 = st.columns(2)
    with col1:
        filter_status = st.selectbox(
            "Filter by Status",
            ["All"] + GLASSWARE_STATUS_OPTIONS,
            key="glassware_status_filter"
        )
    with col2:
        search_term = st.text_input("Search Glassware", key="glassware_search")
    
    # Apply filters
    filtered_df = df.copy()
    if filter_status != "All":
        filtered_df = filtered_df[filtered_df["status"] == filter_status]
    if search_term:
        filtered_df = filtered_df[
            filtered_df["item"].str.contains(search_term, case=False) |
            filtered_df["id"].str.contains(search_term, case=False)
        ]

    # Enhanced row styling function
    def highlight_glassware_issues(row):
        styles = [''] * len(row)
        
        # Highlight zero quantity items
        if row.get('current', 0) <= 0:
            styles = ['background-color: #FFCCCC'] * len(row)  # Light red for zero quantity
        
        # Highlight different statuses with different colors
        status = row.get('status', '').lower()
        if 'broken' in status:
            styles = ['background-color: #FF9999'] * len(row)  # Stronger red for broken
        elif 'chipped' in status:
            styles = ['background-color: #FFCC99'] * len(row)  # Orange for chipped
        elif 'missing' in status:
            styles = ['background-color: #FFFF99'] * len(row)  # Yellow for missing
        
        return styles
    
    # Display the styled dataframe first for better visibility
    st.dataframe(
        filtered_df.style.apply(highlight_glassware_issues, axis=1),
        use_container_width=True,
        height=600,
        column_order=['id', 'item', 'category', 'total_quantity', 'broken_quantity', 'current', 'status', 'location']
    )
    
    # Enhanced data editor with broken quantity tracking
    edited_df = st.data_editor(
        filtered_df,
        num_rows="dynamic",
        use_container_width=True,
        disabled=["id", "current"],
        column_config={
            "status": st.column_config.SelectboxColumn(
                "Status",
                options=GLASSWARE_STATUS_OPTIONS,
                required=True
            ),
            "total_quantity": st.column_config.NumberColumn(
                "Total Qty",
                min_value=0,
                step=1
            ),
            "broken_quantity": st.column_config.NumberColumn(
                "Broken Qty",
                min_value=0,
                step=1
            ),
            "current": st.column_config.NumberColumn(
                "Usable Qty",
                help="Automatically calculated as Total - Broken"
            )
        }
    )
    
    # Update status based on broken quantity (but don't update current in the DataFrame)
    if "broken_quantity" in edited_df.columns:
        edited_df["status"] = np.where(
            edited_df["broken_quantity"] > 0,
            "Broken",
            "OK"
        )
    
    # Update status based on broken quantity (but don't update current in the DataFrame)
    if "broken_quantity" in edited_df.columns:
        # Only update status if it's not already set to a specific condition
        edited_df["status"] = np.where(
            (edited_df["broken_quantity"] > 0) & (edited_df["status"] == "OK"),
            "Broken",
            edited_df["status"]
        )

    if st.button("Save Glassware Changes"):
        # Remove the 'current' column before saving since it's generated by the database
        if 'current' in edited_df.columns:
            edited_df = edited_df.drop(columns=['current'])
        
        save_to_db("glassware", edited_df)
        st.session_state.lab_inventory["glassware"] = pd.read_sql(
            text("SELECT * FROM glassware"), 
            get_conn()
        )
        st.success("Changes saved successfully!")
        st.rerun()

def manage_equipment():
    """Equipment inventory management with non-working items tracking"""
    df = st.session_state.lab_inventory["equipment"].copy()
    
    # Convert string dates to datetime objects if they exist
    if 'last_calibration' in df.columns:
        df['last_calibration'] = pd.to_datetime(
            df['last_calibration'], 
            format='%Y-%m-%d', 
            errors='coerce'
        )
    
    with st.expander("➕ Add New Equipment", expanded=False):
        with st.form("add_equipment_form"):
            cols = st.columns([1, 2, 1])
            with cols[0]:
                new_id = generate_unique_id("equipment")
                st.text_input("ID", value=new_id, disabled=True)
                item_name = st.text_input("Item Name", key="new_equip_name")
                category = st.selectbox("Category", INVENTORY_CATEGORIES["equipment"])
            with cols[1]:
                total_quantity = st.number_input("Total Quantity", min_value=1, step=1, value=1)
                non_working_quantity = st.number_input("Non-Working Quantity", min_value=0, step=1, value=0, 
                                                     max_value=total_quantity)
                working_quantity = total_quantity - non_working_quantity
                st.text_input("Working Quantity", value=working_quantity, disabled=True)
                unit = st.selectbox("Unit", ["units", "sets"])
            with cols[2]:
                location = st.text_input("Storage Location")
                status = st.selectbox("Status", ["Working", "Needs service", "In repair", "Broken", "Out for calibration"])
                last_calibration = st.date_input(
                    "Last Calibration Date",
                    value=datetime.now().date() - timedelta(days=30),
                    format="YYYY-MM-DD"
                )
                comment = st.text_input("Comments")
            
            if st.form_submit_button("Add Equipment"):
                existing_items = df['item'].str.lower().tolist()
                if item_name.lower() in existing_items:
                    st.error(f"An equipment item with the name '{item_name}' already exists in Inventory!")
                else:
                    new_item = {
                        "id": new_id,
                        "item": item_name,
                        "category": category,
                        "total_quantity": total_quantity,
                        "non_working_quantity": non_working_quantity,
                        "unit": unit,
                        "location": location,
                        "status": status,
                        "last_calibration": last_calibration.strftime('%Y-%m-%d'),
                        "comment": comment
                    }
                    
                    # Double-check ID doesn't exist
                    conn = get_conn()
                    if conn:
                        existing_id = conn.execute(
                            text("SELECT id FROM equipment WHERE id = :id"),
                            {"id": new_id}
                        ).fetchone()
                        
                        if existing_id:
                            st.error(f"ID {new_id} already exists! Generating new ID...")
                            new_id = generate_unique_id("equipment")
                            new_item["id"] = new_id
                    
                    result = insert_item("equipment", new_item)
                    if result:
                        st.session_state.lab_inventory["equipment"] = pd.read_sql(
                            text("SELECT * FROM equipment"), 
                            get_conn()
                        )
                        st.success(f"Added {item_name} to inventory")
                        st.rerun()
    
    st.subheader("Current Equipment Inventory")
    
    # Enhanced row styling function for equipment
    def highlight_equipment_issues(row):
        styles = [''] * len(row)
        
        # Highlight different statuses with different colors
        status = row.get('status', '').lower()
        non_working = row.get('non_working_quantity', 0)
        total = row.get('total_quantity', 1)
        
        if non_working == total:  # All items non-working
            styles = ['background-color: #FF9999'] * len(row)  # Strong red
        elif non_working > 0:     # Some items non-working
            styles = ['background-color: #FFCC99'] * len(row)  # Orange
        elif 'broken' in status:
            styles = ['background-color: #FF9999'] * len(row)  # Strong red
        elif 'repair' in status or 'service' in status or 'calibration' in status:
            styles = ['background-color: #FFFF99'] * len(row)  # Yellow
        
        return styles
    
    # Display the styled dataframe first for better visibility
    st.dataframe(
        df.style.apply(highlight_equipment_issues, axis=1),
        use_container_width=True,
        height=600,
        column_order=['id', 'item', 'category', 'total_quantity', 'non_working_quantity', 'status', 'last_calibration']
    )
    
    # Enhanced data editor with non-working quantity tracking
    edited_df = st.data_editor(
        df,
        num_rows="dynamic",
        use_container_width=True,
        disabled=["id"],
        column_config={
            "status": st.column_config.SelectboxColumn(
                "Status",
                options=["Working", "Needs service", "In repair", "Broken", "Out for calibration"],
                required=True,
                help="Current operational status of the equipment"
            ),
            "total_quantity": st.column_config.NumberColumn(
                "Total Qty",
                min_value=1,
                step=1,
                help="Total number of this equipment including non-working units"
            ),
            "non_working_quantity": st.column_config.NumberColumn(
                "Non-Working Qty",
                min_value=0,
                step=1,
                help="Number of non-functional units"
            ),
            "last_calibration": st.column_config.DateColumn(
                "Last Calibration",
                format="YYYY-MM-DD",
                min_value=date(2000, 1, 1),
                help="Select last calibration date from calendar"
            )
        }
    )
    
    if st.button("Save Equipment Changes"):
        # Convert datetime back to strings before saving
        if 'last_calibration' in edited_df.columns:
            edited_df['last_calibration'] = edited_df['last_calibration'].dt.strftime('%Y-%m-%d')
        
        save_to_db("equipment", edited_df)
        st.session_state.lab_inventory["equipment"] = pd.read_sql(
            text("SELECT * FROM equipment"), 
            get_conn()
        )
        st.success("Changes saved successfully!")
        st.rerun()

# ======================
# Enhanced Visualization Functions
# ======================

def display_inventory_dashboard():
    """Display comprehensive inventory dashboard with visualizations and inventory tables"""
    st.title("📊 Inventory Dashboard")
    
    # Summary statistics
    st.subheader("Inventory Summary")
    cols = st.columns(4)
    
    total_items = 0
    low_stock_items = 0
    expired_items = 0
    damaged_items = 0
    
    for category in ["chemicals", "glassware", "equipment"]:
        stats = calculate_category_stats(category)
        total_items += stats['total']
        low_stock_items += stats.get('low', 0)
        expired_items += stats.get('expired', 0)
        damaged_items += stats.get('damaged', 0)
    
    with cols[0]:
        st.metric("Total Items", total_items)
    with cols[1]:
        st.metric("Low Stock Items", low_stock_items, delta=f"-{low_stock_items} items")
    with cols[2]:
        st.metric("Expired Items", expired_items, delta_color="inverse")
    with cols[3]:
        st.metric("Damaged Items", damaged_items, delta_color="inverse")
    
    # Tabs for different sections
    tab1, tab2, tab3, tab4 = st.tabs(["Visualizations", "All Inventory", "Chemicals", "Glassware & Equipment"])
    
    with tab1:
        # Existing visualization content
        st.subheader("Inventory Status Overview")
        
        # Combine data from all categories
        combined_data = []
        for category in ["chemicals", "glassware", "equipment"]:
            df = st.session_state.lab_inventory[category].copy()
            df['category_type'] = category.capitalize()
            combined_data.append(df)
        
        combined_df = pd.concat(combined_data, ignore_index=True)
        
        # Status distribution pie chart
        if not combined_df.empty:
            col1, col2 = st.columns(2)
            with col1:
                fig1 = px.pie(
                    combined_df, 
                    names='status', 
                    title='Overall Status Distribution',
                    color='status',
                    color_discrete_map={
                        'OK': '#2ecc71',
                        'Low': '#f39c12',
                        'Critical': '#e74c3c',
                        'Expired': '#c0392b',
                        'Damaged': '#c0392b',
                        'Needs service': '#f39c12',
                        'In repair': '#f39c12'
                    }
                )
                st.plotly_chart(fig1, use_container_width=True)
            
            with col2:
                fig2 = px.bar(
                    combined_df.groupby(['category_type', 'status']).size().reset_index(name='count'),
                    x='category_type',
                    y='count',
                    color='status',
                    title='Status by Category',
                    labels={'category_type': 'Category', 'count': 'Count'},
                    color_discrete_map={
                        'OK': '#2ecc71',
                        'Low': '#f39c12',
                        'Critical': '#e74c3c',
                        'Expired': '#c0392b',
                        'Damaged': '#c0392b',
                        'Needs service': '#f39c12',
                        'In repair': '#f39c12'
                    }
                )
                st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("No inventory data available for visualization.")
        
        # Stock level and expiry analysis
        st.subheader("Detailed Analysis")
        
        if not st.session_state.lab_inventory["chemicals"].empty:
            chem_df = st.session_state.lab_inventory["chemicals"].copy()
            
            # Stock level analysis
            if 'current' in chem_df.columns and 'minimum' in chem_df.columns:
                chem_df = chem_df.dropna(subset=['current', 'minimum'])
                chem_df['percent_of_min'] = (chem_df['current'] / chem_df['minimum']) * 100
                
                fig3 = px.bar(
                    chem_df.sort_values('percent_of_min'),
                    x='item',
                    y='percent_of_min',
                    title='Stock Levels (% of Minimum)',
                    labels={'percent_of_min': '% of Minimum Stock'},
                    color='percent_of_min',
                    color_continuous_scale='RdYlGn_r'
                )
                st.plotly_chart(fig3, use_container_width=True)
            
            # Expiry analysis
            if 'expiry' in chem_df.columns:
                chem_df['expiry'] = pd.to_datetime(chem_df['expiry'], errors='coerce')
                chem_df = chem_df.dropna(subset=['expiry'])
                chem_df['days_until_expiry'] = (chem_df['expiry'] - pd.to_datetime('today')).dt.days
                
                fig4 = px.bar(
                    chem_df.sort_values('days_until_expiry'),
                    x='item',
                    y='days_until_expiry',
                    title='Days Until Expiry',
                    color='days_until_expiry',
                    color_continuous_scale='YlOrRd_r'
                )
                st.plotly_chart(fig4, use_container_width=True)
    
    with tab2:
        # All Inventory Table
        st.subheader("Complete Inventory List")
        
        combined_data = []
        for category in ["chemicals", "glassware", "equipment"]:
            df = st.session_state.lab_inventory[category].copy()
            df['Type'] = category.capitalize()
            combined_data.append(df)
        
        combined_df = pd.concat(combined_data, ignore_index=True)
        
        if not combined_df.empty:
            def color_code_row(row):
                colors = [''] * len(row)
                status = row.get('status', '')
                
                if status == 'Low':
                    colors = ['background-color: #FFF3CD'] * len(row)
                elif status in ['Expired', 'Critical', 'Damaged']:
                    colors = ['background-color: #F8D7DA'] * len(row)
                elif status in ['Needs service', 'In repair']:
                    colors = ['background-color: #FFF3CD'] * len(row)
                
                return colors
            
            st.dataframe(
                combined_df.style.apply(color_code_row, axis=1),
                use_container_width=True,
                height=600,
                column_order=['Type', 'id', 'item', 'category', 'current', 'minimum', 'status', 'expiry']
            )
        else:
            st.info("No inventory data available")
    
    with tab3:
        # Chemicals Table
        st.subheader("Chemical Inventory")
        chem_df = st.session_state.lab_inventory["chemicals"].copy()
        
        if not chem_df.empty:
            def color_code_chemicals(row):
                colors = [''] * len(row)
                status = row.get('status', '')
                
                if status == 'Low':
                    colors = ['background-color: #FFF3CD'] * len(row)
                elif status in ['Expired', 'Critical']:
                    colors = ['background-color: #F8D7DA'] * len(row)
                
                return colors
            
            st.dataframe(
                chem_df.style.apply(color_code_chemicals, axis=1),
                use_container_width=True,
                height=600,
                column_order=['id', 'item', 'category', 'current', 'minimum', 'monthly', 'Months Stock', 'status', 'expiry']
            )
        else:
            st.info("No chemical inventory data available")
    
    with tab4:
        # Glassware and Equipment Tables
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Glassware Inventory")
            glass_df = st.session_state.lab_inventory["glassware"].copy()
            
            if not glass_df.empty:
                def color_code_glassware(row):
                    colors = [''] * len(row)
                    if row.get('status', '') == 'Damaged':
                        colors = ['background-color: #F8D7DA'] * len(row)
                    return colors
                
                st.dataframe(
                    glass_df.style.apply(color_code_glassware, axis=1),
                    use_container_width=True,
                    height=600
                )
            else:
                st.info("No glassware inventory data available")
        
        with col2:
            st.subheader("Equipment Inventory")
            equip_df = st.session_state.lab_inventory["equipment"].copy()
            
            if not equip_df.empty:
                def color_code_equipment(row):
                    colors = [''] * len(row)
                    if row.get('status', '') in ['Needs service', 'In repair']:
                        colors = ['background-color: #FFF3CD'] * len(row)
                    return colors
                
                st.dataframe(
                    equip_df.style.apply(color_code_equipment, axis=1),
                    use_container_width=True,
                    height=600
                )
            else:
                st.info("No equipment inventory data available")
                
# ======================
# Enhanced Reporting Functions
# ======================

def generate_comprehensive_report():
    """Generate a comprehensive inventory report with visualizations"""
    # Create a buffer for the report
    buffer = BytesIO()
    
    # Create a document
    doc = Document()
    doc.add_heading('Laboratory Inventory Comprehensive Report', 0)
    
    # Add report metadata
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Report generated on: {current_time}")
    doc.add_paragraph(f"Generated by: {st.session_state.username}")
    
    # Add summary statistics
    doc.add_heading('Executive Summary', level=1)
    
    # Calculate summary stats
    total_items = 0
    low_stock = 0
    expired = 0
    damaged = 0
    issues = 0
    
    for category in ["chemicals", "glassware", "equipment"]:
        stats = calculate_category_stats(category)
        total_items += stats['total']
        low_stock += stats.get('low', 0)
        expired += stats.get('expired', 0)
        damaged += stats.get('damaged', 0)
        issues += stats.get('issues', 0)
    
    # Add summary table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Value'
    
    # Data rows
    table.rows[1].cells[0].text = 'Total Inventory Items'
    table.rows[1].cells[1].text = str(total_items)
    
    table.rows[2].cells[0].text = 'Items Low on Stock'
    table.rows[2].cells[1].text = str(low_stock)
    
    table.rows[3].cells[0].text = 'Expired Items'
    table.rows[3].cells[1].text = str(expired)
    
    table.rows[4].cells[0].text = 'Items Needing Attention'
    table.rows[4].cells[1].text = str(damaged + issues)
    
    # Add detailed sections for each category
    for category in ["chemicals", "glassware", "equipment"]:
        doc.add_heading(f"{category.capitalize()} Inventory Details", level=1)
        
        df = st.session_state.lab_inventory[category].copy()
        
        # Add summary for the category
        stats = calculate_category_stats(category)
        
        # Create a table for category summary
        cat_table = doc.add_table(rows=len(stats)+1, cols=2)
        cat_table.style = 'Table Grid'
        
        # Header row
        cat_table.rows[0].cells[0].text = 'Metric'
        cat_table.rows[0].cells[1].text = 'Value'
        
        # Data rows
        for i, (key, value) in enumerate(stats.items(), start=1):
            cat_table.rows[i].cells[0].text = key.replace('_', ' ').title()
            cat_table.rows[i].cells[1].text = str(value)
    
    # Save the document to the buffer
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def show_reporting_section():
    """Display the enhanced reporting interface"""
    st.header("📑 Inventory Reporting")
    
    with st.expander("📄 Generate Comprehensive Report", expanded=True):
        report_type = st.selectbox(
            "Select Report Type", 
            ["Full Inventory Report", "Chemicals Report", "Glassware Report", "Equipment Report"]
        )
        
        if st.button("Generate Report"):
            with st.spinner("Generating report..."):
                report_buffer = generate_comprehensive_report()
                st.success("Report generated successfully!")
                
                st.download_button(
                    label="Download Report",
                    data=report_buffer,
                    file_name=f"lab_inventory_report_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Email report option
                with st.expander("📧 Email Report"):
                    receiver_email = st.text_input("Recipient Email")
                    email_subject = st.text_input("Subject", value=f"Lab Inventory Report - {datetime.now().strftime('%Y-%m-%d')}")
                    email_body = st.text_area("Email Body", value="Please find attached the latest laboratory inventory report.")
                    
                    if st.button("Send Email"):
                        if not receiver_email:
                            st.error("Please enter a recipient email address")
                        else:
                            if send_email_report(receiver_email, email_subject, email_body, report_buffer):
                                st.success("Email sent successfully!")
                            else:
                                st.error("Failed to send email")
    
    with st.expander("📊 Quick Visualizations", expanded=False):
        st.subheader("Quick Inventory Visualizations")
        
        # Chemicals visualization
        if not st.session_state.lab_inventory["chemicals"].empty:
            st.markdown("#### Chemicals Status")
            chem_df = st.session_state.lab_inventory["chemicals"].copy()
            
            # Status distribution
            fig1 = px.pie(chem_df, names='status', title='Chemical Status Distribution')
            st.plotly_chart(fig1, use_container_width=True)
            
            # Months of stock histogram
            if 'Months Stock' in chem_df.columns:
                fig2 = px.histogram(
                    chem_df, 
                    x='Months Stock', 
                    title='Months of Stock Remaining',
                    nbins=20
                )
                st.plotly_chart(fig2, use_container_width=True)
        
        # Equipment visualization
        if not st.session_state.lab_inventory["equipment"].empty:
            st.markdown("#### Equipment Status")
            equip_df = st.session_state.lab_inventory["equipment"].copy()
            
            fig3 = px.pie(equip_df, names='status', title='Equipment Status Distribution')
            st.plotly_chart(fig3, use_container_width=True)

# ======================
# Updated Main Page
# ======================

def display_lab_inventory_page():
    """Main inventory management interface with enhanced visualization"""
    st.title("🧪 Laboratory Inventory Management System")
    
    if 'lab_inventory' not in st.session_state:
        initialize_inventory_data()
    
    # Navigation tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📊 Dashboard", 
        "🧪 Chemicals", 
        "🧫 Glassware", 
        "⚙️ Equipment", 
        "📑 Reports"
    ])
    
    with tab1:
        display_inventory_dashboard()
    
    with tab2:
        manage_chemicals()
    
    with tab3:
        manage_glassware()
    
    with tab4:
        manage_equipment()
    
    with tab5:
        show_reporting_section()

if __name__ == "__main__":
    try:
        conn = get_conn()
        if conn:
            try:
                init_inventory_tables()
                st.session_state.lab_inventory = load_from_db()
                
                if all(df.empty for df in st.session_state.lab_inventory.values()):
                    load_sample_data()
                    for category, df in st.session_state.lab_inventory.items():
                        save_to_db(category, df)
                    st.session_state.lab_inventory = load_from_db()
                
                check_restock_status()
                display_lab_inventory_page()
            finally:
                conn.close()
    except Exception as e:
        st.error(f"Application error: {e}")