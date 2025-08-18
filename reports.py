import streamlit as st
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import tempfile
import os
import base64
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime as dt


def generate_report(report_data, report_type, start_date, end_date, comments=None):
    """
    Generate a professional Word document report with visualizations
    
    Args:
        report_data: DataFrame with the processed report data
        report_type: Type of report being generated
        start_date: Start date of report period
        end_date: End date of report period
        comments: Optional comments to include
        
    Returns:
        Path to the generated Word document
    """
    # Create document
    doc = Document()
    
    # Add title with formatting
    title = doc.add_heading(f'{report_type} Quality Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add report period
    period = doc.add_paragraph()
    period.add_run('Report Period: ').bold = True
    period.add_run(f"{start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}")
    
    # Add company logo if available
    try:
        doc.add_picture('logo.png', width=Inches(1.5))
    except:
        pass
    
    # Add comments section if provided
    if comments:
        doc.add_heading('Comments', level=2)
        doc.add_paragraph(comments)
    
   # Check if we have data to process
    if report_data.empty:
        doc.add_paragraph("No data available for this report period.")
    else:
        # Process each report section - handle both 'Section' and 'Report Section' column names
        section_column = None
        for possible_column in ['Section', 'Report Section']:
            if possible_column in report_data.columns:
                section_column = possible_column
                break
    
    # Process each report section
    if section_column:
        sections = report_data[section_column].unique()
    else:
        # Fallback if column was renamed
        sections = ['All Data']  # Single section for all data
    
    for section in sections:
        doc.add_heading(section, level=1)
        if section_column in report_data.columns:
            section_data = report_data[report_data['Section'] == section]
        else:
            section_data = report_data  # Use all data if no section column
        
        # Create table with available columns
            available_columns = [col for col in ['Metric', 'Value'] if col in section_data.columns]
            
            if available_columns:
                table = doc.add_table(rows=1, cols=len(available_columns))
                table.style = 'Light Grid'
        
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            # Add data rows
            for _, row in section_data.iterrows():
                row_cells = table.add_row().cells
                for i, col in enumerate(available_columns):
                            row_cells[i].text = str(row.get(col, ''))
            
            # Add visualizations for key sections if data is available
                if 'BRIX' in section and 'Value' in section_data.columns:
                    try:
                        add_brix_chart(doc, section_data)
                    except Exception as e:
                        doc.add_paragraph(f"Could not generate BRIX chart: {str(e)}")
                elif 'Torque' in section and 'Value' in section_data.columns:
                    try:
                        add_torque_chart(doc, section_data)
                    except Exception as e:
                        doc.add_paragraph(f"Could not generate Torque chart: {str(e)}")
    
    # Add footer
    doc.add_paragraph().add_run(f"Report generated on {dt.datetime.now().strftime('%B %d, %Y at %H:%M')}").italic = True
    
    # Save to temporary file
    temp_dir = tempfile.mkdtemp()
    report_path = os.path.join(temp_dir, f'Quality_Report_{start_date.strftime("%Y%m%d")}_to_{end_date.strftime("%Y%m%d")}.docx')
    doc.save(report_path)
    
    return report_path

def add_brix_chart(doc, data):
    """Add BRIX visualization to document"""
    metrics = data['Metric'].tolist()
    values = [float(v.split()[0]) for v in data['Value']]
    
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(metrics, values, color='#4e79a7')
    ax.set_title('BRIX Measurements', pad=20)
    ax.set_ylabel('BRIX Value')
    
    # Add value labels
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.2f}',
                ha='center', va='bottom')
    
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    # Add to document
    add_figure_to_doc(doc, fig, 'BRIX Measurement Results')
    plt.close(fig)

def add_torque_chart(doc, data):
    """Add Torque visualization to document"""
    metrics = data['Metric'].tolist()
    values = [float(v.split()[0]) if '%' not in v else float(v.split('%')[0]) 
              for v in data['Value']]
    
    fig, ax = plt.subplots(figsize=(10, 5))
    colors = ['#4e79a7' if 'Average' in m else 
              '#f28e2b' if 'Out of Range' in m else 
              '#e15759' for m in metrics]
    
    bars = ax.bar(metrics, values, color=colors)
    ax.set_title('Torque Performance', pad=20)
    ax.set_ylabel('Value (Nm)' if 'Average' in metrics[0] else 'Percentage')
    
    # Add value labels
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.1f}{"%" if "%" in str(data["Value"].iloc[0]) else ""}',
                ha='center', va='bottom')
    
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    # Add to document
    add_figure_to_doc(doc, fig, 'Torque Measurement Results')
    plt.close(fig)

def add_figure_to_doc(doc, fig, caption):
    """Helper to add matplotlib figures to Word doc"""
    temp_dir = tempfile.mkdtemp()
    img_path = os.path.join(temp_dir, 'temp_img.png')
    fig.savefig(img_path, dpi=300, bbox_inches='tight')
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(img_path)
    os.rmdir(temp_dir)

def download_report(report_data, filename_prefix):
    """
    Create a download link for the Word report
    
    Args:
        report_data: DataFrame with report data
        filename_prefix: Prefix for the download filename
    """
    # Parse the filename prefix to extract components
    parts = filename_prefix.split('_')
    if len(parts) < 4:
        st.error("Invalid filename format. Expected: ReportType_YYYY-MM-DD_to_YYYY-MM-DD")
        return
    
    report_type = parts[0]
    start_date_str = parts[1]
    end_date_str = parts[3]
    
    try:
        # Convert string dates to datetime.date objects
        start_date = dt.datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = dt.datetime.strptime(end_date_str, '%Y-%m-%d').date()

        # Generate the Word document
        report_path = generate_report(
            report_data=report_data,
            report_type=report_type,
            start_date=start_date,
            end_date=end_date
        )
    
        # Create download link
        with open(report_path, 'rb') as f:
            report_bytes = f.read()
        
        b64 = base64.b64encode(report_bytes).decode()
        filename = f"{filename_prefix}.docx"
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Word Report</a>'
        
        st.markdown(href, unsafe_allow_html=True)
        
        # Clean up
        os.remove(report_path)
        os.rmdir(os.path.dirname(report_path))

    except ValueError as e:
            st.error(f"Invalid date format in filename: {e}")
    except Exception as e:
            st.error(f"Error generating report: {e}")